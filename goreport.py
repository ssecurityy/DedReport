#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
This is the GoReport class. GoReport handles everything from connecting to the target Gophish
server to pulling campaign information and reporting the results.
"""

try:
    # 3rd Party Libraries
    from gophish import Gophish
except:
    print("[!] Could not import the Gophish library! Make sure it is installed.\n\
Run: `python3 -m pip intall gophish`\n\
Test it by running `python3` and then, in the \
Python prompt, typing `from gophish import Gophish`.")
    exit()

# Standard Libraries
import configparser
import os.path
import sys
from collections import Counter

# 3rd Party Libraries
import requests
import xlsxwriter
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor
from requests.packages.urllib3.exceptions import InsecureRequestWarning
from user_agents import parse
import multiprocessing
import re
import datetime
from collections import defaultdict

requests.packages.urllib3.disable_warnings(InsecureRequestWarning)


class Goreport(object):
    """
    This class uses the Gophish library to create a new Gophish API connection
    and queries Gophish for information and results related to the specified
    campaign ID(s).
    """
    # Name of the config file -- default is ``gophish.config``
    goreport_config_file = "gophish.config"
    verbose = False

    # Variables for holding Gophish models
    results = None
    campaign = None
    timeline = None

    # Variables for holding campaign information
    cam_id = None
    cam_url = None
    cam_name = None
    cam_status = None
    launch_date = None
    created_date = None
    cam_page_name = None
    cam_smtp_host = None
    completed_date = None
    cam_redirect_url = None
    cam_from_address = None
    cam_subject_line = None
    cam_template_name = None
    cam_capturing_passwords = None
    cam_capturing_credentials = None

    # Variables and lists for tracking event numbers
    total_sent = 0
    total_opened = 0
    total_targets = 0
    total_clicked = 0
    total_reported = 0
    total_submitted = 0
    total_unique_opened = 0
    total_unique_clicked = 0
    total_unique_reported = 0
    total_unique_submitted = 0
    targets_opened = []
    targets_clicked = []
    targets_reported = []
    targets_submitted = []

    # Lists and dicts for holding prepared report data
    campaign_results_summary = []

    # Lists for holding totals for statistics
    browsers = []
    locations = []
    ip_addresses = []
    ip_and_location = {}
    operating_systems = []

    # Output options
    report_format = None
    output_word_report = None
    output_xlsx_report = None
    xlsx_header_bg_color = "#0085CA"
    xlsx_header_font_color = "#FFFFFF"

    def __init__(self, report_format, config_file, google, verbose):
        """
        Initiate the connection to the Gophish server with the provided host, port,
        and API key and prepare to use the external APIs.
        """
        try:
            # Check if an alternate config file was provided
            if config_file:
                self.goreport_config_file = config_file
            # Open the config file to make sure it exists and is readable
            config = configparser.ConfigParser()
            config.read(self.goreport_config_file)
        except Exception as e:
            print(f"[!] Could not open {self.goreport_config_file} -- make sure it exists and is readable.")
            print(f"L.. Details: {e}")
            sys.exit()

        try:
            # Read in the values from the config file
            GP_HOST = self.config_section_map(config, 'Gophish')['gp_host']
            API_KEY = self.config_section_map(config, 'Gophish')['api_key']
        except Exception as e:
            print("[!] There was a problem reading values from the gophish.config file!")
            print(f"L.. Details: {e}")
            sys.exit()

        try:
            # Read in the values from the config file
            self.IPINFO_TOKEN = self.config_section_map(config, 'ipinfo.io')['ipinfo_token']
            if not self.IPINFO_TOKEN:
                self.IPINFO_TOKEN = None
        except Exception as e:
            self.IPINFO_TOKEN = None
            print("[!] No ipinfo.io API token was found in the config. GoReport will not lookup IP addresses with ipinfo.io for additional location data.")
            print(f"L.. Details: {e}")

        try:
            # Read in the values from the config file
            self.GEOLOCATE_TOKEN = self.config_section_map(config, 'Google')['geolocate_key']
            if not self.GEOLOCATE_TOKEN:
                self.GEOLOCATE_TOKEN = None
        except Exception as e:
            self.GEOLOCATE_TOKEN = None
            if google:
                print("[!] No Google Maps API token was found in the config so GoReport will ignore the `--google` flag.")
                print(f"L.. Details: {e}")

        # Set command line options for the GoReport object
        self.google = google
        self.verbose = verbose
        self.report_format = report_format
        # Connect to the Gophish API
        # NOTE: This step succeeds even with a bad API key, so the true test is fetching an ID
        print(f"[+] Connecting to Gophish at {GP_HOST}")
        print(f"L.. The API Authorization endpoint is: {GP_HOST}/api/campaigns/?api_key={API_KEY}")
        self.api = Gophish(API_KEY, host=GP_HOST, verify=False)

    def run(self, id_list, combine_reports, set_complete_status):
        """Run everything to process the target campaign."""
        # Output some feedback for user options
        if combine_reports:
            print("[+] Campaign results will be combined into a single report.")
        if set_complete_status:
            print('[+] Campaign statuses will be set to "Complete" after processing the results.')
        try:
            # Create the list of campaign IDs
            temp_id = []
            # Handle a mixed set of ranges and comma-separated IDs
            if "-" and "," in id_list:
                temp = id_list.split(",")
                for x in temp:
                    if "-" in x:
                        lower = x.split("-")[0]
                        upper = x.split("-")[1]
                        for y in range(int(lower), int(upper) + 1):
                            temp_id.append(str(y))
                    else:
                        temp_id.append(x)
            # Process IDs provided as one or more ranges
            elif "-" in id_list:
                lower = id_list.split("-")[0]
                upper = id_list.split("-")[1]
                for y in range(int(lower), int(upper) + 1):
                    temp_id.append(str(y))
            # Process single or only comma-separated IDs
            else:
                temp_id = id_list.split(",")
            id_list = temp_id
        except Exception as e:
            print("[!] Could not interpret your provided campaign IDs. \
Ensure the IDs are provided as comma-separated integers or interger ranges, e.g. 5,50-55,71.")
            print(f"L.. Details: {e}")
            sys.exit()
        # Begin processing the campaign IDs by removing any duplicates
        try:
            # Get length of user-provided list
            initial_len = len(id_list)
            # Remove duplicate IDs and sort IDs as integers
            id_list = sorted(set(id_list), key=int)
            # Get length of unique, sorted list
            unique_len = len(id_list)
        except Exception as e:
            temp = []
            for id in id_list:
                try:
                    int(id)
                except:
                    temp.append(id)
            print(f"[!] There are {len(temp)} invalid campaign ID(s), i.e. not an integer.")
            print(f"L.. Offending IDs: {','.join(temp)}")
            print(f"L.. Details: {e}")
            sys.exit()
        print(f"[+] A total of {initial_len} campaign IDs have been provided for processing.")
        # If the lengths are different, then GoReport removed one or more dupes
        if initial_len != unique_len:
            dupes = initial_len - unique_len
            print(f"L.. GoReport found {dupes} duplicate campaign IDs, so those have been trimmed.")
        # Provide  list of all IDs that will be processed
        print(f"[+] GoReport will process the following campaign IDs: {','.join(id_list)}")
        # If --combine is used with just one ID it can break reporting, so we catch that here
        if len(id_list) == 1 and combine_reports:
            combine_reports = False
        # Go through each campaign ID and get the results
        campaign_counter = 1
        for CAM_ID in id_list:
            print(f"[+] Now fetching results for Campaign ID {CAM_ID} ({campaign_counter}/{len(id_list)}).")
            try:
                # Request the details for the provided campaign ID
                self.campaign = self.api.campaigns.get(campaign_id=CAM_ID)
            except Exception as e:
                print(f"[!] There was a problem fetching this campaign {CAM_ID}'s details. Make sure your URL and API key are correct. Check HTTP vs HTTPS!")
                print(f"L.. Details: {e}")
            try:
                try:
                    # Check to see if a success message was returned with a message
                    # Possible reasons: campaign ID doesn't exist or problem with host/API key
                    if self.campaign.success is False:
                        print(f"[!] Failed to get results for campaign ID {CAM_ID}")
                        print(f"L.. Details: {self.campaign.message}")
                        # We can't let an error with an ID stop reporting, so check if this was the last ID
                        if CAM_ID == id_list[-1] and combine_reports:
                            self.generate_report()
                # If self.campaign.success does not exist then we were successful
                except:
                    print("[+] Success!")
                    # Collect campaign details and process data
                    self.collect_all_campaign_info(combine_reports)
                    self.process_timeline_events(combine_reports)
                    self.process_results(combine_reports)
                    # If the --complete flag was set, now set campaign status to Complete
                    if set_complete_status:
                        print(f"[+] Setting campaign ID {CAM_ID}'s status to Complete.")
                        try:
                            set_complete = self.api.campaigns.complete(CAM_ID)
                            try:
                                if set_complete.success is False:
                                    print(f"[!] Failed to set campaign status for ID {CAM_ID}.")
                                    print(f"L.. Details: {set_complete.message}")
                            # If set_complete.success does not exist then we were successful
                            except:
                                pass
                        except Exception as e:
                            print(f"[!] Failed to set campaign status for ID {CAM_ID}.")
                            print(f"L.. Details: {e}")
                    # Check if this is the last campaign ID in the list
                    # If this is the last ID and combined reports is on, generate the report
                    if CAM_ID == id_list[-1] and combine_reports:
                        self.generate_report()
                    # Otherwise, if we are not combining reports, generate the reports
                    elif combine_reports is False:
                        self.generate_report()
                    campaign_counter += 1
            except Exception as e:
                print(f"[!] There was a problem processing campaign ID {CAM_ID}!")
                print(f"L.. Details: {e}")
                sys.exit()

    def lookup_ip(self, ip):
        """Lookup the provided IP address with ipinfo.io for location data.

        Example Result:
            {'ip': '52.44.93.197',
            'hostname': 'ec2-52-44-93-197.compute-1.amazonaws.com',
            'city': 'Beaumont',
            'region': 'Texas',
            'country': 'US',
            'loc': '30.0866,-94.1274',
            'postal': '77702',
            'phone': '409',
            'org': 'AS14618 Amazon.com, Inc.'}
        """
        ipinfo_url = f"https://ipinfo.io/{ip}?token={self.IPINFO_TOKEN}"
        try:
            r = requests.get(ipinfo_url)
            return r.json()
        except Exception as e:
            print(f"[!] Failed to lookup `{ip}` with ipinfo.io.")
            print(f"L.. Details: {e}")
            return None

    def get_google_location_data(self, lat, lon):
        """Use Google's Maps API to collect location info for the provided latitude and longitude.

        Google returns a bunch of JSON with a variety of location data. This function returns
        Google's pre-formatted `formatted_address` key for a human-readable address.
        """
        google_maps_url = f"https://maps.googleapis.com/maps/api/geocode/json?latlng={lat},{lon}&sensor=false&key={self.GEOLOCATE_TOKEN}"
        r = requests.get(google_maps_url)
        maps_json = r.json()
        if r.ok:
            try:
                if "error_message" in maps_json:
                    print(f"[!] Google Maps returned an error so using Gophish coordinates. Error: {maps_json['error_message']}")
                    return f"{lat}, {lon}"
                first_result = maps_json['results'][0]
                if "formatted_address" in first_result:
                    return first_result["formatted_address"]
                # In case that key is ever unavailable try to assemble an address
                else:
                    components = first_result['address_components']
                    country = town = None
                    for c in components:
                        if "country" in c['types']:
                            country = c['long_name']
                        if "locality" in c['types']:
                            town = c['long_name']
                        if "administrative_area_level_1" in c['types']:
                            state = c['long_name']
                    return f"{town}, {state}, {country}"
            except Exception as e:
                print("[!] Failed to parse Google Maps API results so using Gophish coordinates.")
                print(f"L.. Error: {e}")
                return f"{lat}, {lon}"
        else:
            print(f"[!] Failed to contact the Google Maps API so using Gophish coordinates. Status code: {r.status_code}")
            return f"{lat}, {lon}"

    def geolocate(self, target, ipaddr, google=False):
        """Attempt to get location data for the provided target and event. Will use ipinfo.io if an
        API key is configured. Otherwise the Gophish latitude and longitude coordinates will be
        returned. If `google` is set to True this function will try to match the coordinates to a
        location using the Google Maps API.

        Returns a string: City, Region, Country
        """
        if ipaddr in self.ip_and_location:
            return self.ip_and_location[ipaddr]
        else:
            if self.IPINFO_TOKEN:
                # location_json = self.lookup_ip(event.details['browser']['address'])
                location_json = self.lookup_ip(ipaddr)
                if location_json:
                    city = region = country = "Unknown"
                    if "city" in location_json:
                        if location_json['city']:
                            city = location_json['city']
                    if "region" in location_json:
                        if location_json['region']:
                            region = location_json['region']
                    if "country" in location_json:
                        if location_json['country']:
                            country = location_json['country']
                    location = f"{city}, {region}, {country}"
                else:
                    location = f"{target.latitude}, {target.longitude}"
            elif google:
                if self.GEOLOCATE_TOKEN:
                    location = self.get_google_location_data(target.latitude, target.longitude)
                else:
                    location = f"{target.latitude}, {target.longitude}"
            else:
                location = f"{target.latitude}, {target.longitude}"
            self.locations.append(location)
            self.ip_and_location[ipaddr] = location
            return location

    def compare_ip_addresses(self, target_ip, browser_ip, verbose):
        """Compare the IP addresses of the target to that of an event. The goal: Looking for a
        mismatch that might identify some sort of interesting event. This might indicate an
        email was forwarded, a VPN was switched on/off, or maybe the target is at home.
        """
        if target_ip == browser_ip:
            return target_ip
        else:
            # We have an IP mismatch -- hard to tell why this might be.
            if verbose:
                print(f"[*] Event: This target's ({target_ip}) URL was clicked from a browser at {browser_ip}.")
            # This is an IP address not included in the results model, so we add it to our list here
            self.ip_addresses.append(browser_ip)
            return browser_ip

    def get_basic_campaign_info(self):
        """"Helper function to collect a campaign's basic details. This includes campaign name,
        status, template, and other details that are not the campaign's results.

        This keeps these calls in one place for tidiness and easier management.
        """
        self.cam_name = self.campaign.name
        self.cam_status = self.campaign.status
        self.created_date = self.campaign.created_date
        self.launch_date = self.campaign.launch_date
        self.completed_date = self.campaign.completed_date
        self.cam_url = self.campaign.url

        # Collect SMTP information
        self.smtp = self.campaign.smtp
        self.cam_from_address = self.smtp.from_address
        self.cam_smtp_host = self.smtp.host

        # Collect the template information
        self.template = self.campaign.template
        self.cam_subject_line = self.template.subject
        self.cam_template_name = self.template.name
        self.cam_template_attachments = self.template.attachments
        if self.cam_template_attachments == []:
            self.cam_template_attachments = "None Used"

        # Collect the landing page information
        self.page = self.campaign.page
        self.cam_page_name = self.page.name
        self.cam_redirect_url = self.page.redirect_url
        if self.cam_redirect_url == "":
            self.cam_redirect_url = "Not Used"
        self.cam_capturing_passwords = self.page.capture_passwords
        self.cam_capturing_credentials = self.page.capture_credentials

    def collect_all_campaign_info(self, combine_reports):
        """Collect the campaign's details and set values for each of the variables."""
        # Collect the basic campaign details
        try:
            # Begin by checking if the ID is valid
            self.cam_id = self.campaign.id
            if combine_reports and self.cam_name is None:
                print(f"[+] Reports will be combined -- setting name, dates, and URL based on campaign ID {self.cam_id}.")
                self.get_basic_campaign_info()
            elif combine_reports is False:
                self.get_basic_campaign_info()
            # Collect the results and timeline lists
            if self.results is None:
                self.results = self.campaign.results
                self.timeline = self.campaign.timeline
            elif combine_reports:
                self.results += self.campaign.results
                self.timeline += self.campaign.timeline
            else:
                self.results = self.campaign.results
                self.timeline = self.campaign.timeline
        except:
            print(f"[!] Looks like campaign ID {self.cam_id} does not exist! Skipping it...")

    def process_results(self, combine_reports):
        """Process the results model to collect basic data, like total targets and event details.
        This should be run after the process_timeline_events() function which creates the
        targets_* lists.

        The results model can provide:
        first_name, last_name, email, position, and IP address
        """
        # Total length of results gives us the total number of targets
        if combine_reports and self.total_targets is None:
            self.total_targets = len(self.campaign.results)
        elif combine_reports:
            self.total_targets += len(self.campaign.results)
        else:
            # Not combining, so reset counters
            self.total_unique_opened = 0
            self.total_unique_clicked = 0
            self.total_unique_reported = 0
            self.total_unique_submitted = 0
            # Reports will not be combined, so reset tracking between reports
            self.total_targets = len(self.campaign.results)
            self.ip_addresses = []
            self.campaign_results_summary = []
        # Go through all results and extract data for statistics
        for target in self.campaign.results:
            temp_dict = {}
            # Log the IP address for additional statistics later
            if not target.ip == "":
                self.ip_addresses.append(target.ip)
                self.geolocate(target, target.ip, self.google)
            # Add all of the recipient's details and results to the temp dictionary
            temp_dict["email"] = target.email
            temp_dict["fname"] = target.first_name
            temp_dict["lname"] = target.last_name
            position = "None Provided"
            if target.position:
                position = target.position
            temp_dict["position"] = position
            temp_dict["ip_address"] = target.ip
            # Check if this target was recorded as viewing the email (tracking image)
            if target.email in self.targets_opened:
                temp_dict["opened"] = True
                self.total_unique_opened += 1
            else:
                temp_dict["opened"] = False
            # Check if this target clicked the link
            if target.email in self.targets_clicked:
                temp_dict["clicked"] = True
                self.total_unique_clicked += 1
                # Incremement the total number of opens for this target if they clicked
                # but did not display the tracking image in the email
                if target.email not in self.targets_opened:
                    self.total_unique_opened += 1
            else:
                temp_dict["clicked"] = False
            # Check if this target submitted data
            if target.email in self.targets_submitted:
                temp_dict["submitted"] = True
                self.total_unique_submitted += 1
            else:
                temp_dict["submitted"] = False
            # Check if this target reported the email
            if target.email in self.targets_reported:
                temp_dict["reported"] = True
                self.total_unique_reported += 1
            else:
                temp_dict["reported"] = False
            # Append the temp dictionary to the event summary list
            self.campaign_results_summary.append(temp_dict)

    def process_timeline_events(self, combine_reports):
        """Process the timeline model to collect basic data, like total clicks, and get detailed
        event data for recipients.

        The timeline model contains all events that occurred during the campaign.
        """
        # Create counters for enumeration
        sent_counter = 0
        click_counter = 0
        opened_counter = 0
        reported_counter = 0
        submitted_counter = 0

        # Reset target lists
        self.targets_opened = []
        self.targets_clicked = []
        self.targets_reported = []
        self.targets_submitted = []
        # Run through all events and count each of the four basic events
        for event in self.campaign.timeline:
            if event.message == "Email Sent":
                sent_counter += 1
            elif event.message == "Email Opened":
                opened_counter += 1
                self.targets_opened.append(event.email)
            elif event.message == "Clicked Link":
                click_counter += 1
                self.targets_clicked.append(event.email)
            elif event.message == "Submitted Data":
                submitted_counter += 1
                self.targets_submitted.append(event.email)
            elif event.message == "Email Reported":
                reported_counter += 1
                self.targets_reported.append(event.email)
        # Assign the counter values to our tracking lists
        if combine_reports:
            # Append, +=, totals if combining reports
            self.total_sent += sent_counter
            self.total_opened += opened_counter
            self.total_clicked += click_counter
            self.total_reported += reported_counter
            self.total_submitted += submitted_counter
        else:
            # Set tracking variables to current counter values for non-combined reports
            self.total_sent = sent_counter
            self.total_opened = opened_counter
            self.total_clicked = click_counter
            self.total_reported = reported_counter
            self.total_submitted = submitted_counter

    def generate_report(self):
        """Determines which type of report generate and the calls the appropriate reporting
        functions.
        """
        if self.report_format == "excel":
            print("[+] Building the report -- you selected a Excel/xlsx report.")
            self.output_xlsx_report = self._build_output_xlsx_file_name()
            self.write_xlsx_report()
        elif self.report_format == "word":
            print("[+] Building the report -- you selected a Word/docx report.")
            print("[+] Looking for the template.docx to be used for the Word report.")
            if os.path.isfile("template.docx"):
                print("[+] Template was found -- proceeding with report generation...")
                print("L.. Word reports can take a while if you had a lot of recipients.")
                self.output_word_report = self._build_output_word_file_name()
                self.write_word_report()
            else:
                print("[!] Could not find the template document! Make sure 'template.docx' is in the GoReport directory.")
                sys.exit()
        elif self.report_format == "html":
            print("[+] Building the report -- you selected an HTML report.")
            self.output_html_report = self._build_output_html_file_name()
            self.write_html_report()
        elif self.report_format == "quick":
            print("[+] Quick report stats:")
            self.get_quick_stats()

    def get_quick_stats(self):
        """Present quick stats for the campaign. Just basic numbers and some details."""
        print()
        print(self.cam_name)
        print(f"Status:\t\t{self.cam_status}")
        print(f"Created:\t{self.created_date.split('T')[1].split('.')[0]} on {self.created_date.split('T')[0]}")
        print(f"Started:\t{self.launch_date.split('T')[1].split('.')[0]} on {self.launch_date.split('T')[0]}")
        if self.cam_status == "Completed":
            print(f"Completed:\t{self.completed_date.split('T')[1].split('.')[0]} on {self.completed_date.split('T')[0]}")
        print()
        print(f"Total Targets:\t{self.total_targets}")
        print(f"Emails Sent:\t{self.total_sent}")
        print(f"IPs Seen:\t{len(self.ip_addresses)}")
        print()
        print(f"Total Opened Events:\t\t{self.total_opened}")
        print(f"Total Click Events:\t\t{self.total_clicked}")
        print(f"Total Submitted Data Events:\t{self.total_submitted}")
        print()
        print(f"Individuals Who Opened:\t\t\t{self.total_unique_opened}")
        print(f"Individuals Who Clicked:\t\t{self.total_unique_clicked}")
        print(f"Individuals Who Entered Data:\t\t{self.total_unique_submitted}")
        print(f"Individuals Who Reported the Email:\t{self.total_unique_reported}")

    def _build_output_xlsx_file_name(self):
        """Create the xlsx report name."""
        safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        xlsx_report = f"Gophish Results for {safe_name}.xlsx"
        return xlsx_report

    def _build_output_word_file_name(self):
        """Create the docx report name."""
        safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        word_report = f"Gophish Results for {safe_name}.docx"
        return word_report

    def _build_output_html_file_name(self):
        """Create the HTML report name."""
        safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        html_report = f"Gophish Results for {safe_name}.html"
        return html_report

    def _set_word_column_width(self, column, width):
        """Custom function for quickly and easily setting the width of a table's column in the Word
        docx output.

        This option is missing from the basic Python-docx library.
        """
        for cell in column.cells:
            cell.width = width

    def write_xlsx_report(self):
        """Assemble and output the xlsx file report.

        Throughout this function, results are assembled by adding commas and then adding to a
        results string, i.e. 'result_A' and then 'result_A' += ',result_B'. This is so the
        result can be written to the csv file and have the different pieces end up in the correct
        columns.
        """
        goreport_xlsx = xlsxwriter.Workbook(self.output_xlsx_report)
        # Bold format
        bold_format = goreport_xlsx.add_format({'bold': True})
        bold_format.set_text_wrap()
        bold_format.set_align('vcenter')
        # Centered format
        center_format = goreport_xlsx.add_format()
        center_format.set_text_wrap()
        center_format.set_align('vcenter')
        center_format.set_align('center')
        # Header format
        header_format = goreport_xlsx.add_format({'bold': True})
        header_format.set_text_wrap()
        header_format.set_align('vcenter')
        header_format.set_bg_color(self.xlsx_header_bg_color)
        header_format.set_font_color(self.xlsx_header_font_color)
        # Number cells
        num_format = goreport_xlsx.add_format()
        num_format.set_align('center')
        # Boolean cells - True
        true_format = goreport_xlsx.add_format({'bold': True})
        true_format.set_text_wrap()
        true_format.set_align('vcenter')
        true_format.set_font_color("#9C0006")
        true_format.set_bg_color("#FFC7CE")
        # Boolean cells - True
        false_format = goreport_xlsx.add_format()
        false_format.set_text_wrap()
        false_format.set_align('vcenter')
        false_format.set_font_color("#006100")
        false_format.set_bg_color("#C6EFCE")
        # Remaining cells
        wrap_format = goreport_xlsx.add_format()
        wrap_format.set_text_wrap()
        wrap_format.set_align('vcenter')

        worksheet = goreport_xlsx.add_worksheet("Overview")
        col = 0
        row = 0

        worksheet.set_column(0, 10, 62)

        worksheet.write(row, col, "Campaign Results For:", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_name}", wrap_format)
        row += 1
        worksheet.write(row, col, "Status", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_status}", wrap_format)
        row += 1
        worksheet.write(row, col, "Created", bold_format)
        worksheet.write(row, col + 1, f"{self.created_date}", wrap_format)
        row += 1
        worksheet.write(row, col, "Started", bold_format)
        worksheet.write(row, col + 1, f"{self.launch_date}", wrap_format)
        row += 1
        if self.cam_status == "Completed":
            worksheet.write(row, col, "Completed", bold_format)
            worksheet.write(row, col + 1, f"{self.completed_date}", wrap_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Campaign Details", bold_format)
        row += 1
        worksheet.write(row, col, "From", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_from_address}", wrap_format)
        row += 1
        worksheet.write(row, col, "Subject", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_subject_line}", wrap_format)
        row += 1
        worksheet.write(row, col, "Phish URL", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_url}", wrap_format)
        row += 1
        worksheet.write(row, col, "Redirect URL", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_redirect_url}", wrap_format)
        row += 1
        worksheet.write(row, col, "Attachment(s)", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_template_attachments}", wrap_format)
        row += 1
        worksheet.write(row, col, "Captured Passwords", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_capturing_credentials}", wrap_format)
        row += 1
        worksheet.write(row, col, "Stored Passwords", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_capturing_passwords}", wrap_format)
        row += 1

        worksheet.write(row, col, "")
        row += 1

        # Write a high level summary for stats
        worksheet.write(row, col, "High Level Results", bold_format)
        row += 1
        worksheet.write(row, col, "Total Targets", bold_format)
        worksheet.write(row, col + 1, self.total_targets, num_format)
        row += 1

        worksheet.write(row, col, "The following totals indicate how many events of each type Gophish recorded:", wrap_format)
        row += 1
        worksheet.write(row, col, "Total Opened Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_opened, num_format)
        row += 1
        worksheet.write(row, col, "Total Clicked Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_clicked, num_format)
        row += 1
        worksheet.write(row, col, "Total Submitted Data Events", bold_format)
        worksheet.write(row, col + 1, "", wrap_format)
        row += 1
        worksheet.write(row, col, "Total Report Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_reported, num_format)
        row += 1

        worksheet.write(row, col, "The following totals indicate how many targets participated in each event type:", wrap_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Opened", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_opened, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Clicked", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_clicked, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Submitted Data", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_submitted, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Reported", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_reported, num_format)
        row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet = goreport_xlsx.add_worksheet("Summary")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 20)

        worksheet.write(row, col, "Summary of Events", bold_format)
        row += 1

        header_col = 0
        headers = ["Email Address", "Open", "Click", "Creds", "Report", "OS", "Browser"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1

        # Sort campaign summary by each dict's email entry and then create results table
        target_counter = 0
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        for target in ordered_results:
            color_info = self.get_action_color(target['opened'], target['clicked'], target['submitted'], target['reported'])
            row_format = goreport_xlsx.add_format({'bg_color': color_info['excel']})
            worksheet.write(row, col, target['email'], row_format)
            worksheet.write_boolean(row, col + 1, target['opened'], row_format)
            worksheet.write_boolean(row, col + 2, target['clicked'], row_format)
            worksheet.write_boolean(row, col + 3, target['submitted'], row_format)
            worksheet.write_boolean(row, col + 4, target['reported'], row_format)
            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 5, browser_details, row_format)
                        worksheet.write(row, col + 6, os_details, row_format)
            else:
                worksheet.write(row, col + 5, "N/A", row_format)
                worksheet.write(row, col + 6, "N/A", row_format)
            row += 1
            target_counter += 1
            print(f"[+] Created row for {target_counter} of {self.total_targets}.")

        print("[+] Finished writing events summary...")
        print("[+] Detailed results analysis is next and will take some time if you had a lot of targets...")
        # End of the event summary and beginning of the detailed results

        worksheet = goreport_xlsx.add_worksheet("Event Details")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 40)

        worksheet.write(row, col, "Detailed Analysis", bold_format)
        row += 1

        target_counter = 0
        for target in self.results:
            # Only create a Detailed Analysis section for targets with clicks
            if target.email in self.targets_clicked:
                position = ""
                if target.position:
                    position = f"({target.position})"
                worksheet.write(row, col, f"{target.first_name} {target.last_name} {position}", bold_format)
                row += 1
                worksheet.write(row, col, target.email, wrap_format)
                row += 1
                # Go through all events to find events for this target
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        # Parse the timestamp into separate date and time variables
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]
                        # Record the email sent date and time in the report
                        worksheet.write(row, col, f"Sent on {sent_date.replace(',', '')} at {sent_time}", wrap_format)
                        row += 1

                    if event.message == "Email Opened" and event.email == target.email:
                        # Record the email preview date and time in the report
                        temp = event.time.split('T')
                        worksheet.write(row, col, f"Email Preview at {temp[0]} {temp[1].split('.')[0]}", wrap_format)
                        row += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        worksheet.write(row, col, "Email Link Clicked", bold_format)
                        row += 1

                        header_col = 0
                        headers = ["Time", "IP", "Location", "Browser", "Operating System"]
                        for header in headers:
                            worksheet.write(row, header_col, header, header_format)
                            header_col += 1
                        row += 1

                        temp = event.time.split('T')
                        worksheet.write(row, col, f"{temp[0]} {temp[1].split('.')[0]}", wrap_format)

                        # Check if browser IP matches the target's IP and record result
                        ip_comparison = self.compare_ip_addresses(target.ip,
                                                                  event.details['browser']['address'],
                                                                  self.verbose)
                        worksheet.write(row, col + 1, f"{ip_comparison}", wrap_format)

                        # Parse the location data
                        loc = self.geolocate(target, event.details['browser']['address'], self.google)
                        worksheet.write(row, col + 2, loc, wrap_format)

                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        worksheet.write(row, col + 3, browser_details, wrap_format)
                        self.browsers.append(browser_details)

                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 4, os_details, wrap_format)
                        self.operating_systems.append(os_details)
                        row += 1

                    if event.message == "Submitted Data" and event.email == target.email:
                        # Now we have events for submitted data. A few notes on this:
                        #   1. There is no expectation of a Submit event without a Clicked Link event
                        #   2. Assuming that, the following process does NOT flag IP mismatches
                        #      or add to the list of seen locations, OSs, IPs, or browsers.
                        worksheet.write(row, col, "Submitted Data Captured", bold_format)
                        row += 1

                        header_col = 0
                        headers = ["Time", "IP", "Location", "Browser", "Operating System", "Data Captured"]
                        for header in headers:
                            worksheet.write(row, header_col, header, header_format)
                            header_col += 1
                        row += 1

                        temp = event.time.split('T')
                        worksheet.write(row, col, f"{temp[0]} {temp[1].split('.')[0]}", wrap_format)

                        worksheet.write(row, col + 1, f"{event.details['browser']['address']}", wrap_format)

                        loc = self.geolocate(target, event.details['browser']['address'], self.google)
                        worksheet.write(row, col + 2, loc, wrap_format)

                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        worksheet.write(row, col + 3, browser_details, wrap_format)

                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 4, os_details, wrap_format)

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += f"{key}:{str(value).strip('[').strip(']')}"
                        worksheet.write(row, col + 5, submitted_data, wrap_format)
                        row += 1

                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
            else:
                # This target had no clicked or submitted events so move on to next
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
                continue
            worksheet.write(row, col, "")
            row += 1

        print("[+] Finished writing detailed analysis...")

        worksheet = goreport_xlsx.add_worksheet("Stats")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 35)

        worksheet.write(row, col, "Recorded Browsers Based on User-Agents:", bold_format)
        row += 1

        header_col = 0
        headers = ["Browser", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_browsers = Counter(self.browsers)
        for key, value in counted_browsers.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Record OS From Browser User-Agents:", bold_format)
        row += 1
        header_col = 0
        headers = ["Operating System", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_os = Counter(self.operating_systems)
        for key, value in counted_os.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Recorded Locations from IPs:", bold_format)
        row += 1
        header_col = 0
        headers = ["Locations", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Recorded IPs:", bold_format)
        row += 1
        header_col = 0
        headers = ["IP Address", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "Recorded IPs and Locations:", bold_format)
        row += 1
        header_col = 0
        headers = ["IP Address", "Location"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        for key, value in self.ip_and_location.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write(row, col + 1, f"{value}", wrap_format)
            row += 1

        goreport_xlsx.close()
        print(f"[+] Done! Check '{self.output_xlsx_report}' for your results.")

    def write_word_report(self):
        """Assemble and output the Word docx file report."""
        import time
        from collections import defaultdict
        start_time = time.time()
        d = Document("template.docx")
        styles = d.styles
        # Branding update
        for style in d.styles:
            if hasattr(style, 'name') and style.name == 'GoReport':
                style.name = 'Dedsec Technologies'
        # Create a custom styles for table cells
        _ = styles.add_style("Cell Text", WD_STYLE_TYPE.CHARACTER)
        cell_text = d.styles["Cell Text"]
        cell_text_font = cell_text.font
        cell_text_font.name = "Calibri"
        cell_text_font.size = Pt(12)
        cell_text_font.bold = True
        cell_text_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        _ = styles.add_style("Cell Text Hit", WD_STYLE_TYPE.CHARACTER)
        cell_text_hit = d.styles["Cell Text Hit"]
        cell_text_hit_font = cell_text_hit.font
        cell_text_hit_font.name = "Calibri"
        cell_text_hit_font.size = Pt(12)
        cell_text_hit_font.bold = True
        cell_text_hit_font.color.rgb = RGBColor(0x00, 0x96, 0x00)

        _ = styles.add_style("Cell Text Miss", WD_STYLE_TYPE.CHARACTER)
        cell_text_miss = d.styles["Cell Text Miss"]
        cell_text_miss_font = cell_text_miss.font
        cell_text_miss_font.name = "Calibri"
        cell_text_miss_font.size = Pt(12)
        cell_text_miss_font.bold = True
        cell_text_miss_font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Write a campaign summary at the top of the report
        d.add_heading("Executive Summary", 1)
        p = d.add_paragraph()
        run = p.add_run(f"Campaign Results For: {self.cam_name}")
        run.bold = True
        # Runs are basically "runs" of text and must be aligned like we want
        # them aligned in the report -- thus they are pushed left
        if self.cam_status == "Completed":
            completed_status = f"Completed:\t{self.completed_date.split('T')[1].split('.')[0]} on {self.completed_date.split('T')[0]}"
        else:
            completed_status = "Still Active"
        p.add_run(f"""
Status: {self.cam_status}
Created: {self.created_date.split('T')[1].split('.')[0]} on {self.created_date.split('T')[0]}
Started: {self.launch_date.split('T')[1].split('.')[0]} on {self.launch_date.split('T')[0]}
Completed: {completed_status}

""")
        if self.cam_status == "Completed":
            print()

        # Write the campaign details -- email details and template settings
        run = p.add_run("Campaign Details")
        run.bold = True
        p.add_run(f"""
From: {self.cam_from_address}
Subject: {self.cam_subject_line}
Phish URL: {self.cam_url}
Redirect URL: {self.cam_redirect_url}
Attachment(s): {self.cam_template_attachments}
Captured Credentials: {self.cam_capturing_credentials}
Stored Passwords: {self.cam_capturing_passwords}

""")

        # Write a high level summary for stats
        run = p.add_run("High Level Results")
        run.bold = True
        p.add_run(f"""
Total Targets: {self.total_targets}

The following totals indicate how many events of each type Gophish recorded:
Total Open Events: {self.total_opened} <span style='font-size:0.7em;color:#888;'>({(self.total_opened/self.total_targets*100):.1f}%)</span>
Total Click Events: {self.total_clicked} <span style='font-size:0.7em;color:#888;'>({(self.total_clicked/self.total_targets*100):.1f}%)</span>
Total Report Events: {self.total_reported} <span style='font-size:0.7em;color:#888;'>({(self.total_reported/self.total_targets*100):.1f}%)</span>
Total Submitted Data Events: {self.total_submitted} <span style='font-size:0.7em;color:#888;'>({(self.total_submitted/self.total_targets*100):.1f}%)</span>

The following totals indicate how many targets participated in each event type:
Individuals Who Opened: {self.total_unique_opened}
Individuals Who Clicked: {self.total_unique_clicked}
Individuals Who Reported: {self.total_unique_reported}
Individuals Who Submitted: {self.total_unique_submitted}

""")
        d.add_page_break()

        print("[Dedsec Technologies] Finished writing high level summary...")
        # End of the campaign summary and beginning of the event summary
        d.add_heading("Summary of Events", 1)
        d.add_paragraph("The following table summarizes who opened and clicked on emails sent in this campaign.")

        # Create a table to hold the event summary results
        table = d.add_table(rows=len(self.campaign_results_summary) + 1, cols=7, style="Dedsec Technologies")

        header0 = table.cell(0, 0)
        header0.text = ""
        header0.paragraphs[0].add_run("Email Address", "Cell Text").bold = True

        header1 = table.cell(0, 1)
        header1.text = ""
        header1.paragraphs[0].add_run("Open", "Cell Text").bold = True

        header2 = table.cell(0, 2)
        header2.text = ""
        header2.paragraphs[0].add_run("Click", "Cell Text").bold = True

        header3 = table.cell(0, 3)
        header3.text = ""
        header3.paragraphs[0].add_run("Data", "Cell Text").bold = True

        header4 = table.cell(0, 4)
        header4.text = ""
        header4.paragraphs[0].add_run("Report", "Cell Text").bold = True

        header5 = table.cell(0, 5)
        header5.text = ""
        header5.paragraphs[0].add_run("OS", "Cell Text").bold = True

        header6 = table.cell(0, 6)
        header6.text = ""
        header6.paragraphs[0].add_run("Browser", "Cell Text").bold = True

        # Sort campaign summary by each dict's email entry and then create results table
        target_counter = 0
        counter = 1
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        # Build all row data in memory first
        table_rows = []
        for target in ordered_results:
            row = [target['email']]
            row.append(u'\u2713' if target['opened'] else u'\u2718')
            row.append(u'\u2713' if target['clicked'] else u'\u2718')
            row.append(u'\u2713' if target['submitted'] else u'\u2718')
            row.append(u'\u2713' if target['reported'] else u'\u2718')
            if target['email'] in self.targets_clicked:
                os_details = "N/A"
                browser_details = "N/A"
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        break
                row.append(os_details)
                row.append(browser_details)
            else:
                row.append("N/A")
                row.append("N/A")
            table_rows.append(row)
        # Now add all rows to the docx table in one go
        for idx, row in enumerate(table_rows, start=1):
            for col_idx, value in enumerate(row):
                cell = table.cell(idx, col_idx)
                if col_idx == 0:
                    cell.text = value
                elif col_idx in [1,2,3,4]:
                    if value == u'\u2713':
                        cell.paragraphs[0].add_run(value, "Cell Text Hit")
                    else:
                        cell.paragraphs[0].add_run(value, "Cell Text Miss")
                else:
                    cell.text = value
            print(f"[+] Created table entry for {idx} of {self.total_targets}.")

        d.add_page_break()

        # End of the event summary and beginning of the detailed results
        print("[Dedsec Technologies] Finished writing events summary...")
        print("[Dedsec Technologies] Detailed results analysis is next and may take some time if you had a lot of targets...")
        detailed_start = time.time()
        d.add_heading("Detailed Findings", 1)
        target_counter = multiprocessing.Value('i', 0)
        # Prepare data for multiprocessing
        def process_target(target):
            from docx import Document
            from docx.shared import Cm
            from user_agents import parse
            # Prepare a mini docx for this target
            doc = Document()
            if target.email in self.targets_clicked:
                position = f"({target.position})" if target.position else ""
                doc.add_heading(f"{target.first_name} {target.last_name} {position}", 2)
                doc.add_paragraph(target.email)
                p = doc.add_paragraph()
                email_sent_run = p.add_run()
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]
                        email_sent_run.text = f"Email sent on {sent_date} at {sent_time}"
                    if event.message == "Email Opened" and event.email == target.email:
                        if opened_counter.value == 1:
                            # Create the Email Opened/Previewed table
                            p = doc.add_paragraph()
                            p.style = doc.styles['Normal']
                            run = p.add_run("Email Previews")
                            run.bold = True

                            opened_table = doc.add_table(rows=1, cols=1, style="Dedsec Technologies")
                            opened_table.autofit = True
                            opened_table.allow_autofit = True

                            header1 = opened_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                        # Begin by adding a row to the table and inserting timestamp
                        opened_table.add_row()
                        timestamp = opened_table.cell(opened_counter.value, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]
                        opened_counter.value += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        if clicked_counter.value == 1:
                            # Create the Clicked Link table
                            p = doc.add_paragraph()
                            p.style = doc.styles['Normal']
                            run = p.add_run("Email Link Clicked")
                            run.bold = True

                            clicked_table = doc.add_table(rows=1, cols=5, style="Dedsec Technologies")
                            clicked_table.autofit = True
                            clicked_table.allow_autofit = True

                            header1 = clicked_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                            header2 = clicked_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

                            header3 = clicked_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("Location", "Cell Text").bold = True

                            header4 = clicked_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Browser", "Cell Text").bold = True

                            header5 = clicked_table.cell(0, 4)
                            header5.text = ""
                            header5.paragraphs[0].add_run("Operating System",
                                                          "Cell Text").bold = True

                        clicked_table.add_row()
                        timestamp = clicked_table.cell(clicked_counter.value, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        ip_add = clicked_table.cell(clicked_counter.value, 1)
                        # Check if browser IP matches the target's IP and record result
                        ip_add.text = self.compare_ip_addresses(
                            target.ip, event.details['browser']['address'], self.verbose)

                        # Parse the location data
                        event_location = clicked_table.cell(clicked_counter.value, 2)
                        event_location.text = self.geolocate(target, event.details['browser']['address'], self.google)

                        # Parse the user-agent string for browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        browser = clicked_table.cell(clicked_counter.value, 3)
                        browser.text = browser_details
                        self.browsers.append(browser_details)

                        op_sys = clicked_table.cell(clicked_counter.value, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = os_details
                        self.operating_systems.append(os_details)

                        clicked_counter.value += 1

                    if event.message == "Submitted Data" and event.email == target.email:
                        if submitted_counter.value == 1:
                            # Create the Submitted Data table
                            p = doc.add_paragraph()
                            p.style = doc.styles['Normal']
                            run = p.add_run("Data Captured")
                            run.bold = True

                            submitted_table = doc.add_table(rows=1, cols=6, style="Dedsec Technologies")
                            submitted_table.autofit = True
                            submitted_table.allow_autofit = True

                            header1 = submitted_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                            header2 = submitted_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

                            header3 = submitted_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("Location", "Cell Text").bold = True

                            header4 = submitted_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Browser", "Cell Text").bold = True

                            header5 = submitted_table.cell(0, 4)
                            header5.text = ""
                            header5.paragraphs[0].add_run("Operating System",
                                                          "Cell Text").bold = True

                            header6 = submitted_table.cell(0, 5)
                            header6.text = ""
                            header6.paragraphs[0].add_run("Data Captured",
                                                          "Cell Text").bold = True

                        submitted_table.add_row()
                        timestamp = submitted_table.cell(submitted_counter.value, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        ip_add = submitted_table.cell(submitted_counter.value, 1)
                        ip_add.text = event.details['browser']['address']

                        # Parse the location data
                        event_location = submitted_table.cell(submitted_counter.value, 2)
                        event_location.text = self.geolocate(target, event.details['browser']['address'], self.google)

                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        browser = submitted_table.cell(submitted_counter.value, 3)
                        browser.text = browser_details

                        op_sys = submitted_table.cell(submitted_counter.value, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = f"{os_details}"

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data = submitted_table.cell(submitted_counter.value, 5)
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += f"{key}:{str(value).strip('[').strip(']')}   "
                        data.text = f"{submitted_data}"
                        submitted_counter.value += 1
            return doc
        with multiprocessing.Pool() as pool:
            docs = pool.map(process_target, self.results)
        # Merge all mini-docs into main docx (d)
        for doc in docs:
            for element in doc.element.body:
                d.element.body.append(element)
        print(f"[Dedsec Technologies] Finished writing Detailed Analysis section in {time.time() - detailed_start:.2f} seconds...")
        # End of the detailed results and the beginning of browser, location, and OS stats
        d.add_heading("Statistics", 1)
        p = d.add_paragraph("The following table shows the browsers seen:")
        # Create browser table
        browser_table = d.add_table(rows=1, cols=2, style="Dedsec Technologies")
        self._set_word_column_width(browser_table.columns[0], Cm(7.24))
        self._set_word_column_width(browser_table.columns[1], Cm(3.35))

        header1 = browser_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Browser", "Cell Text").bold = True

        header2 = browser_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the operating systems seen:")

        # Create OS table
        os_table = d.add_table(rows=1, cols=2, style="Dedsec Technologies")
        self._set_word_column_width(os_table.columns[0], Cm(7.24))
        self._set_word_column_width(os_table.columns[1], Cm(3.35))

        header1 = os_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Operating System", "Cell Text").bold = True

        header2 = os_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the locations seen:")

        # Create geo IP table
        location_table = d.add_table(rows=1, cols=2, style="Dedsec Technologies")
        self._set_word_column_width(location_table.columns[0], Cm(7.24))
        self._set_word_column_width(location_table.columns[1], Cm(3.35))

        header1 = location_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Location", "Cell Text").bold = True

        header2 = location_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Visits", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the IP addresses captured:")

        # Create IP address table
        ip_add_table = d.add_table(rows=1, cols=2, style="Dedsec Technologies")
        self._set_word_column_width(ip_add_table.columns[0], Cm(7.24))
        self._set_word_column_width(ip_add_table.columns[1], Cm(3.35))

        header1 = ip_add_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("IP Address", "Cell Text").bold = True

        header2 = ip_add_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the IP addresses matched with geolocation data:")

        # Create IP address and location table
        ip_loc_table = d.add_table(rows=1, cols=2, style="Dedsec Technologies")
        self._set_word_column_width(ip_loc_table.columns[0], Cm(7.24))
        self._set_word_column_width(ip_loc_table.columns[1], Cm(3.35))

        header1 = ip_loc_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("IP Address", "Cell Text").bold = True

        header2 = ip_loc_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Location", "Cell Text").bold = True

        # Counters are used here again to track rows
        counter = 1
        # Counter is used to count all elements in the lists to create a unique list with totals
        counted_browsers = Counter(self.browsers)
        for key, value in counted_browsers.items():
            browser_table.add_row()
            cell = browser_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = browser_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_os = Counter(self.operating_systems)
        for key, value in counted_os.items():
            os_table.add_row()
            cell = os_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = os_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            location_table.add_row()
            cell = location_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = location_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            ip_add_table.add_row()
            cell = ip_add_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = ip_add_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        for key, value in self.ip_and_location.items():
            ip_loc_table.add_row()
            cell = ip_loc_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = ip_loc_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        # Finalize document and save it as the value of output_word_report
        save_start = time.time()
        d.save(f"{self.output_word_report}")
        print(f"[Dedsec Technologies] Done! Check \"{self.output_word_report}\" for your results. Saving took {time.time() - save_start:.2f} seconds.")
        print(f"[Dedsec Technologies] Word report generation took {time.time() - start_time:.2f} seconds.")

    def write_html_report(self):
        """Fully modern, robust HTML report with advanced analytics, UI/UX, and interactivity."""
        import datetime
        import json
        import re
        from collections import Counter, defaultdict
        client_name = input("Enter the client name for branding: ")
        total_targets = self.total_targets
        total_opened = self.total_opened
        total_clicked = self.total_clicked
        total_reported = self.total_reported
        total_submitted = self.total_submitted
        # --- Advanced Executive Summary Analysis ---
        all_passwords = []
        password_users = defaultdict(list)
        password_entropy = []
        domain_counter = Counter()
        geo_counter = Counter()
        device_counter = Counter()
        time_to_action = []
        risk_scores = {}
        for target in self.results:
            domain = target.email.split('@')[-1] if '@' in target.email else ''
            domain_counter[domain] += 1
            # Device/OS/Geo
            for event in self.timeline:
                if hasattr(event, 'email') and event.email == target.email:
                    details = event.details if isinstance(event.details, dict) else {}
                    ua = details.get('browser', {}).get('user-agent', '') if 'browser' in details else ''
                    if ua:
                        from user_agents import parse
                        user_agent = parse(ua)
                        device_counter[user_agent.device.family] += 1
                        geo = details.get('browser', {}).get('address', '')
                        if geo:
                            geo_counter[geo] += 1
        for user in self.results:
            for event in self.timeline:
                if hasattr(event, 'email') and event.email == user.email:
                    details = event.details if isinstance(event.details, dict) else {}
                    payload = details.get('payload', {})
                    if payload:
                        for k, v in payload.items():
                            if 'pass' in k.lower() or 'password' in k.lower():
                                pwd = str(v)
                                all_passwords.append(pwd)
                                password_users[pwd].append(user.email)
                                # Entropy calculation
                                entropy = 0
                                pool = 0
                                if re.search(r'[a-z]', pwd): pool += 26
                                if re.search(r'[A-Z]', pwd): pool += 26
                                if re.search(r'\d', pwd): pool += 10
                                if re.search(r'[^\w]', pwd): pool += 32
                                if pool > 0:
                                    entropy = len(pwd) * (pool.bit_length())
                                password_entropy.append(entropy)
        password_counter = Counter(all_passwords)
        most_common_passwords = password_counter.most_common(10)
        pattern_stats = {'short':0, 'long':0, 'has_digit':0, 'has_special':0, 'reused':0, 'weak':0, 'strong':0}
        seen_passwords = set()
        for pwd in all_passwords:
            if len(pwd) < 8: pattern_stats['short'] += 1
            if len(pwd) >= 12: pattern_stats['long'] += 1
            if re.search(r'\d', pwd): pattern_stats['has_digit'] += 1
            if re.search(r'[^\w]', pwd): pattern_stats['has_special'] += 1
            if pwd in seen_passwords: pattern_stats['reused'] += 1
            seen_passwords.add(pwd)
            # Entropy-based strength
            entropy = 0
            pool = 0
            if re.search(r'[a-z]', pwd): pool += 26
            if re.search(r'[A-Z]', pwd): pool += 26
            if re.search(r'\d', pwd): pool += 10
            if re.search(r'[^\w]', pwd): pool += 32
            if pool > 0:
                entropy = len(pwd) * (pool.bit_length())
            if entropy < 40: pattern_stats['weak'] += 1
            else: pattern_stats['strong'] += 1
        # --- End Advanced Analysis ---
        # --- Prepare user details for summary and recipient sections ---
        user_details = []
        user_map = {}
        for target in self.results:
            if target.email not in user_map:
                user_map[target.email] = {
                    'first_name': target.first_name,
                    'last_name': target.last_name,
                    'position': getattr(target, 'position', ''),
                    'ip': getattr(target, 'ip', ''),
                    'events': [],
                    'submitted': False,
                    'opened': False,
                    'clicked': False,
                    'reported': False,
                    'risk': 0
                }
            for event in self.timeline:
                if hasattr(event, 'email') and event.email == target.email:
                    details = event.details if isinstance(event.details, dict) else {}
                    user_map[target.email]['events'].append({
                        'time': getattr(event, 'time', ''),
                        'message': getattr(event, 'message', ''),
                        'details': details
                    })
                    if event.message == 'Email Opened': user_map[target.email]['opened'] = True
                    if event.message == 'Clicked Link': user_map[target.email]['clicked'] = True
                    if event.message == 'Submitted Data': user_map[target.email]['submitted'] = True
                    if event.message == 'Email Reported': user_map[target.email]['reported'] = True
            # Risk scoring: submitted > clicked > opened > reported
            risk = 0
            if user_map[target.email]['submitted']:
                risk = 3
            elif user_map[target.email]['clicked']:
                risk = 2
            elif user_map[target.email]['opened']:
                risk = 1
            if user_map[target.email]['reported']:
                risk = max(risk-1, 0)
            user_map[target.email]['risk'] = risk
            risk_scores[target.email] = risk
        for email, data in user_map.items():
            user_details.append({
                'email': email,
                'first_name': data['first_name'],
                'last_name': data['last_name'],
                'position': data['position'],
                'ip': data['ip'],
                'events': data['events'],
                'opened': data['opened'],
                'clicked': data['clicked'],
                'submitted': data['submitted'],
                'reported': data['reported'],
                'risk': data['risk']
            })
        # --- HTML/CSS/JS Output ---
        html = []
        html.append(f"<html><head><meta charset='utf-8'><title>{client_name} - Dedsec Technologies Gophish Report</title>")
        html.append("<script src='https://cdn.jsdelivr.net/npm/chart.js'></script>")
        html.append("<style>body{font-family:Segoe UI,Calibri,sans-serif;background:#fff;color:#111;}h1,h2{color:#111;}h1 span{color:#0056b3;}header{background:#fff;color:#111;padding:1.5em 2em 1em 2em;border-radius:12px;margin-bottom:2em;box-shadow:0 2px 12px #0001;display:flex;align-items:center;justify-content:space-between;}header .client{font-size:2em;font-weight:700;letter-spacing:1px;}header .logo{font-size:1.2em;font-weight:400;color:#888;}section{margin-bottom:2em;}table{border-collapse:collapse;width:100%;margin-bottom:2em;}th,td{border:1px solid #eee;padding:10px 8px;text-align:left;}th{background:#222;color:#fff;}tr:nth-child(even){background:#f7f7f7;}footer{margin-top:2em;font-size:1em;color:#888;text-align:center;background:#fff;padding:1em 0;border-top:1px solid #eee;}@keyframes fadein{from{opacity:0;}to{opacity:1;}}.fadein{animation:fadein 1.5s;}.user-table{margin-bottom:3em;border-radius:8px;overflow:hidden;box-shadow:0 2px 8px #0002;}.user-header{background:#0056b3;color:#fff;padding:0.7em 1em;font-size:1.1em;cursor:pointer;user-select:none;}.user-header:hover{background:#003366;}.user-details{display:none;}.user-table.active .user-details{display:table-row-group;}.accordion-section{margin-bottom:2em;border-radius:8px;box-shadow:0 2px 8px #0001;overflow:hidden;} .accordion-header{background:#f0f0f0;color:#111;padding:1em 1.5em;font-size:1.2em;cursor:pointer;user-select:none;border-bottom:1px solid #eee;} .accordion-header:hover{background:#e0e0e0;} .accordion-content{display:none;padding:1.5em;background:#fafafa;} .accordion-section.active .accordion-content{display:block;} .confidential{font-size:1.1em;color:#e84118;font-weight:700;letter-spacing:2px;margin-top:1em;} .badge{display:inline-block;padding:0.2em 0.7em;border-radius:12px;font-size:0.95em;margin-right:0.5em;} .badge-opened{background:#ffe082;color:#111;} .badge-clicked{background:#82b1ff;color:#111;} .badge-submitted{background:#81c784;color:#111;} .badge-reported{background:#ff8a65;color:#111;} .badge-none{background:#eee;color:#888;} .row-anim{transition:background 0.3s;} .row-opened{background:#fffde7;} .row-clicked{background:#e3f2fd;} .row-submitted{background:#e8f5e9;} .row-reported{background:#fff3e0;} .row-none{background:#f7f7f7;} .pagination{display:flex;justify-content:center;align-items:center;margin:1em 0;} .pagination button{margin:0 0.2em;padding:0.4em 1em;border:none;background:#0056b3;color:#fff;border-radius:6px;cursor:pointer;} .pagination button.active{background:#222;} .pagination button:disabled{background:#eee;color:#aaa;cursor:not-allowed;}.risk-badge{padding:0.2em 0.7em;border-radius:12px;font-size:0.95em;margin-right:0.5em;font-weight:700;}.risk-0{background:#b2bec3;color:#222;} .risk-1{background:#ffe082;color:#222;} .risk-2{background:#fab1a0;color:#222;} .risk-3{background:#e84118;color:#fff;}</style>")
        html.append("<script src='https://cdn.jsdelivr.net/npm/chartjs-plugin-datalabels'></script>")
        html.append("<script>function toggleAccordion(id){var s=document.getElementById(id);if(s.classList.contains('active')){s.classList.remove('active');}else{s.classList.add('active');}}function toggleUserTable(id){var t=document.getElementById(id);if(t.classList.contains('active')){t.classList.remove('active');}else{t.classList.add('active');}}</script>")
        html.append("</head><body class='fadein'>")
        html.append(f"<header><div class='client'>{client_name}</div><div class='logo'>Dedsec Technologies Gophish Report</div></header>")
        # Executive Summary Section (Accordion)
        html.append("<div class='accordion-section active' id='exec-summary'><div class='accordion-header' onclick=\"toggleAccordion('exec-summary')\">Executive Summary</div><div class='accordion-content'>")
        html.append(f"<div style='display:flex;flex-wrap:wrap;gap:1.5em 2em;'>")
        # Animated stat cards
        html.append(f"<div class='stat-card' style='flex:1;min-width:180px;'><h3>Total Recipients</h3><div style='font-size:2em;font-weight:700;'>{total_targets}</div></div>")
        html.append(f"<div class='stat-card' style='flex:1;min-width:180px;'><h3>Opened</h3><div style='font-size:2em;font-weight:700;'>{total_opened} <span style='font-size:0.7em;color:#888;'>({(total_opened/total_targets*100):.1f}%)</span></div></div>")
        html.append(f"<div class='stat-card' style='flex:1;min-width:180px;'><h3>Clicked</h3><div style='font-size:2em;font-weight:700;'>{total_clicked} <span style='font-size:0.7em;color:#888;'>({(total_clicked/total_targets*100):.1f}%)</span></div></div>")
        html.append(f"<div class='stat-card' style='flex:1;min-width:180px;'><h3>Submitted</h3><div style='font-size:2em;font-weight:700;'>{total_submitted} <span style='font-size:0.7em;color:#888;'>({(total_submitted/total_targets*100):.1f}%)</span></div></div>")
        html.append(f"<div class='stat-card' style='flex:1;min-width:180px;'><h3>Reported</h3><div style='font-size:2em;font-weight:700;'>{total_reported} <span style='font-size:0.7em;color:#888;'>({(total_reported/total_targets*100):.1f}%)</span></div></div>")
        html.append("</div>")
        # Password stats
        html.append("<h3>Password Submission Stats</h3>")
        html.append(f"<ul><li><b>Total Passwords Submitted:</b> {len(all_passwords)}</li><li><b>Unique Passwords:</b> {len(password_counter)}</li><li><b>Reused Passwords:</b> {pattern_stats['reused']}</li><li><b>Weak Passwords:</b> {pattern_stats['weak']}</li><li><b>Strong Passwords:</b> {pattern_stats['strong']}</li></ul>")
        html.append("<h4>Most Common Passwords</h4><ol>")
        for pwd, count in most_common_passwords:
            html.append(f"<li><b>{pwd}</b> ({count} times, users: {', '.join(password_users[pwd])})</li>")
        html.append("</ol>")
        html.append("<h4>Password Patterns</h4><ul>")
        html.append(f"<li>Short (&lt;8 chars): {pattern_stats['short']}</li>")
        html.append(f"<li>Long (&ge;12 chars): {pattern_stats['long']}</li>")
        html.append(f"<li>Contains Digit: {pattern_stats['has_digit']}</li>")
        html.append(f"<li>Contains Special Char: {pattern_stats['has_special']}</li>")
        html.append("</ul>")
        # Domain stats
        html.append("<h4>Top Domains</h4><ol>")
        for dom, cnt in domain_counter.most_common(5):
            html.append(f"<li>{dom}: {cnt} users</li>")
        html.append("</ol>")
        # Device/Geo stats
        html.append("<h4>Top Devices</h4><ol>")
        for dev, cnt in device_counter.most_common(5):
            html.append(f"<li>{dev}: {cnt} users</li>")
        html.append("</ol>")
        html.append("<h4>Top Locations (IP)</h4><ol>")
        for geo, cnt in geo_counter.most_common(5):
            html.append(f"<li>{geo}: {cnt} users</li>")
        html.append("</ol>")
        # Entropy chart
        html.append("<canvas id='entropyChart' height='60'></canvas>")
        html.append("<canvas id='summaryChart' height='80'></canvas><canvas id='pieChart' height='80'></canvas>")
        html.append("</div></div>")
        # ... rest of the report (campaign details, summary table with filters/pagination/color, recipient details with badges, etc.) ...
        # (The rest of the code would follow the same pattern as previous steps, with advanced UI/UX, filters, pagination, and color coding.)
        # ... existing code ...
        # Campaign Details Section (Accordion)
        html.append("<div class='accordion-section' id='campaign-details'><div class='accordion-header' onclick=\"toggleAccordion('campaign-details')\">Campaign Details</div><div class='accordion-content'>")
        html.append(f"<p><b>From:</b> {self.cam_from_address}<br><b>Subject:</b> {self.cam_subject_line}<br><b>Phish URL:</b> {self.cam_url}<br><b>Redirect URL:</b> {self.cam_redirect_url}<br><b>Attachment(s):</b> {self.cam_template_attachments}<br><b>Captured Credentials:</b> {self.cam_capturing_credentials}<br><b>Stored Passwords:</b> {self.cam_capturing_passwords}</p>")
        html.append("</div></div>")
        # Summary Table Section (Accordion)
        html.append("<div class='accordion-section' id='summary-table'><div class='accordion-header' onclick=\"toggleAccordion('summary-table')\">Summary of Events</div><div class='accordion-content'>")
        # Advanced Filter controls
        html.append("""
        <div style='margin-bottom:1em;display:flex;flex-wrap:wrap;gap:1em;'>
            <label><b>Filter:</b></label>
            <select id='filter-type' onchange='filterTable()'>
                <option value='all'>All</option>
                <option value='opened'>Opened</option>
                <option value='clicked'>Clicked</option>
                <option value='submitted'>Submitted</option>
                <option value='reported'>Reported</option>
                <option value='risk'>Risk</option>
                <option value='os'>OS</option>
                <option value='browser'>Browser</option>
                <option value='domain'>Domain</option>
            </select>
            <input type='text' id='filter-value' placeholder='Enter value...' onkeyup='filterTable()' style='margin-left:0.5em;'>
            <label style='margin-left:1em;'><b>Page Size:</b></label>
            <select id='page-size' onchange='paginateTable("summary-table-main",parseInt(this.value))'>
                <option value='10'>10</option>
                <option value='25'>25</option>
                <option value='50'>50</option>
            </select>
            <button onclick='exportTableToCSV()' style='margin-left:1em;'>Export CSV</button>
        </div>
        """)
        html.append("<table id='summary-table-main'><tr><th>Email Address</th><th>Open</th><th>Click</th><th>Data</th><th>Report</th><th>Risk</th><th>OS</th><th>Browser</th></tr>")
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        for target in ordered_results:
            risk = risk_scores.get(target['email'], 0)
            risk_badge = f"<span class='risk-badge risk-{risk}'>Risk {risk}</span>"
            row_class = 'row-none'
            if target['submitted']:
                row_class = 'row-submitted'
            elif target['clicked']:
                row_class = 'row-clicked'
            elif target['opened']:
                row_class = 'row-opened'
            elif target['reported']:
                row_class = 'row-reported'
            row = f"<tr class='row-anim {row_class}'>"
            row += f"<td>{target['email']}</td>"
            row += f"<td>{'\u2714\ufe0f' if target['opened'] else '\u274c'}</td>"
            row += f"<td>{'\u2714\ufe0f' if target['clicked'] else '\u274c'}</td>"
            row += f"<td>{'\u2714\ufe0f' if target['submitted'] else '\u274c'}</td>"
            row += f"<td>{'\u2714\ufe0f' if target['reported'] else '\u274c'}</td>"
            row += f"<td>{risk_badge}</td>"
            os_details = browser_details = "N/A"
            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        break
            row += f"<td>{os_details}</td><td>{browser_details}</td></tr>"
            html.append(row)
        html.append("</table></div></div>")
        # Add advanced filter and pagination script
        html.append("""
        <script>
        function filterTable() {
            var type = document.getElementById('filter-type').value;
            var value = document.getElementById('filter-value').value.toLowerCase();
            var table = document.getElementById('summary-table-main');
            var rows = table.getElementsByTagName('tr');
            for (var i = 1; i < rows.length; i++) {
                var cells = rows[i].getElementsByTagName('td');
                var show = true;
                if (type !== 'all' && value) {
                    if (type === 'opened' && cells[1].textContent.indexOf('\u2714\ufe0f') === -1) show = false;
                    if (type === 'clicked' && cells[2].textContent.indexOf('\u2714\ufe0f') === -1) show = false;
                    if (type === 'submitted' && cells[3].textContent.indexOf('\u2714\ufe0f') === -1) show = false;
                    if (type === 'reported' && cells[4].textContent.indexOf('\u2714\ufe0f') === -1) show = false;
                    if (type === 'risk' && cells[5].textContent.toLowerCase().indexOf(value) === -1) show = false;
                    if (type === 'os' && cells[6].textContent.toLowerCase().indexOf(value) === -1) show = false;
                    if (type === 'browser' && cells[7].textContent.toLowerCase().indexOf(value) === -1) show = false;
                    if (type === 'domain' && cells[0].textContent.split('@')[1] && cells[0].textContent.split('@')[1].toLowerCase().indexOf(value) === -1) show = false;
                } else if (type !== 'all') {
                    if (type === 'opened' && cells[1].textContent.indexOf('\u2714\ufe0f') === -1) show = false;
                    if (type === 'clicked' && cells[2].textContent.indexOf('\u2714\ufe0f') === -1) show = false;
                    if (type === 'submitted' && cells[3].textContent.indexOf('\u2714\ufe0f') === -1) show = false;
                    if (type === 'reported' && cells[4].textContent.indexOf('\u2714\ufe0f') === -1) show = false;
                } else if (value) {
                    var found = false;
                    for (var j = 0; j < cells.length; j++) {
                        if (cells[j].textContent.toLowerCase().indexOf(value) !== -1) found = true;
                    }
                    show = found;
                }
                rows[i].style.display = show ? '' : 'none';
            }
        }
        function paginateTable(tableId, pageSize){
            var table=document.getElementById(tableId);var rows=table.getElementsByTagName('tr');var page=1;
            function showPage(p){for(var i=1;i<rows.length;i++)rows[i].style.display=(i>=(p-1)*pageSize+1&&i<p*pageSize+1)?'':'none';}
            function createPagination(){var n=Math.ceil((rows.length-1)/pageSize);var pag=document.createElement('div');pag.className='pagination';for(var i=1;i<=n;i++){var btn=document.createElement('button');btn.innerText=i;btn.onclick=(function(i){return function(){page=i;showPage(page);}})(i);pag.appendChild(btn);}table.parentNode.insertBefore(pag,table.nextSibling);}
            showPage(page);createPagination();}
        window.onload=function(){paginateTable('summary-table-main',parseInt(document.getElementById('page-size').value));};
        function exportTableToCSV() {
            var csv = [];
            var rows = document.querySelectorAll('#summary-table-main tr');
            for (var i = 0; i < rows.length; i++) {
                var row = [], cols = rows[i].querySelectorAll('td, th');
                for (var j = 0; j < cols.length; j++) row.push('"' + cols[j].innerText.replace(/"/g, '""') + '"');
                csv.push(row.join(','));
            }
            var csvFile = new Blob([csv.join('\n')], {type: 'text/csv'});
            var downloadLink = document.createElement('a');
            downloadLink.download = 'summary_table.csv';
            downloadLink.href = window.URL.createObjectURL(csvFile);
            downloadLink.style.display = 'none';
            document.body.appendChild(downloadLink);
            downloadLink.click();
        }
        </script>
        """)
        # Per-user details Section (Accordion)
        html.append("<div class='accordion-section' id='user-details'><div class='accordion-header' onclick=\"toggleAccordion('user-details')\">Recipient Details</div><div class='accordion-content'>")
        html.append("<div style='margin-bottom:1em;'><label><b>Filter Recipients:</b></label> <input type='text' id='user-filter' onkeyup='filterUsers()' placeholder='Search by name, email, risk, action...'></div>")
        for idx, user in enumerate(user_details):
            risk_badge = f"<span class='risk-badge risk-{user['risk']}'>Risk {user['risk']}</span>"
            user_row_class = 'row-none'
            if user['submitted']:
                user_row_class = 'row-submitted'
            elif user['clicked']:
                user_row_class = 'row-clicked'
            elif user['opened']:
                user_row_class = 'row-opened'
            elif user['reported']:
                user_row_class = 'row-reported'
            html.append(f"<div class='user-table {user_row_class}' id='user-table-{idx}'><div class='user-header' onclick=\"toggleUserTable('user-table-{idx}')\">{user['first_name']} {user['last_name']} ({user['email']}) - {user['position']} {risk_badge}</div>")
            html.append("<table class='user-details'><tr><th>Time</th><th>Event</th><th>IP</th><th>Browser</th><th>OS</th><th>Payload</th></tr>")
            for event in user['events']:
                details = event['details'] if isinstance(event['details'], dict) else {}
                ip = details.get('browser', {}).get('address', '') if 'browser' in details else ''
                ua = details.get('browser', {}).get('user-agent', '') if 'browser' in details else ''
                browser = os = ''
                if ua:
                    user_agent = parse(ua)
                    browser = user_agent.browser.family + " " + user_agent.browser.version_string
                    os = user_agent.os.family + " " + user_agent.os.version_string
                payload = details.get('payload', {})
                payload_str = "<br>".join([f"<b>{k}:</b> {v}" for k, v in payload.items()]) if payload else ''
                html.append(f"<tr class='event-row'><td>{event['time']}</td><td>{event['message']}</td><td>{ip}</td><td>{browser}</td><td>{os}</td><td>{payload_str}</td></tr>")
            html.append("</table></div>")
        html.append("</div></div>")
        html.append("<script>function filterUsers(){var val=document.getElementById('user-filter').value.toLowerCase();var tables=document.querySelectorAll('.user-table');for(var i=0;i<tables.length;i++){var t=tables[i];var header=t.querySelector('.user-header').textContent.toLowerCase();t.style.display=header.indexOf(val)!==-1?'':'none';}}</script>")
        html.append(f"<footer><span class='confidential'>Confidential</span><br>Generated for {client_name} by <b>Dedsec Technologies LLP</b> &copy; {datetime.datetime.now().year}</footer>")
        # Chart.js script
        html.append(f"<script>const ctx=document.getElementById('summaryChart').getContext('2d');new Chart(ctx,{{type:'bar',data:{{labels:['Total','Opened','Clicked','Reported','Submitted'],datasets:[{{label:'Recipients',data:[{total_targets},{total_opened},{total_clicked},{total_reported},{total_submitted}],backgroundColor:['#0056b3','#0097e6','#e84118','#44bd32','#fbc531']}}]}},options:{{responsive:true,plugins:{{legend:{{display:false}}}}}});const ctx2=document.getElementById('pieChart').getContext('2d');new Chart(ctx2,{{type:'pie',data:{{labels:['Opened','Clicked','Reported','Submitted'],datasets:[{{data:[{total_opened},{total_clicked},{total_reported},{total_submitted}],backgroundColor:['#0097e6','#e84118','#44bd32','#fbc531']}}]}},options:{{responsive:true}}}});</script>")
        # --- Add missing animated/interactive elements and polish ---
        # Add animated stat cards (already present, but add animation CSS)
        html.append("<style>.stat-card{background:#f5f7fa;border-radius:12px;box-shadow:0 2px 8px #0001;padding:1.5em 1em;margin-bottom:1em;transition:transform 0.3s,box-shadow 0.3s;}.stat-card:hover{transform:translateY(-4px) scale(1.03);box-shadow:0 6px 24px #0002;}</style>")
        # Add legend for color coding
        html.append("<div style='margin:1em 0;'><b>Legend:</b> <span class='badge badge-opened'>Opened</span> <span class='badge badge-clicked'>Clicked</span> <span class='badge badge-submitted'>Submitted</span> <span class='badge badge-reported'>Reported</span> <span class='badge badge-none'>None</span></div>")
        # Add time-based trends chart (line chart for events over time)
        html.append("<canvas id='timelineChart' height='60'></canvas>")
        # Prepare timeline data
        timeline_data = defaultdict(lambda: {'opened':0,'clicked':0,'submitted':0,'reported':0})
        for event in self.timeline:
            date = getattr(event, 'time', '')[:10]
            if event.message == 'Email Opened': timeline_data[date]['opened'] += 1
            if event.message == 'Clicked Link': timeline_data[date]['clicked'] += 1
            if event.message == 'Submitted Data': timeline_data[date]['submitted'] += 1
            if event.message == 'Email Reported': timeline_data[date]['reported'] += 1
        dates = sorted(timeline_data.keys())
        opened = [timeline_data[d]['opened'] for d in dates]
        clicked = [timeline_data[d]['clicked'] for d in dates]
        submitted = [timeline_data[d]['submitted'] for d in dates]
        reported = [timeline_data[d]['reported'] for d in dates]
        # Add timeline chart script
        html.append(f"<script>const tctx=document.getElementById('timelineChart').getContext('2d');new Chart(tctx,{{type:'line',data:{{labels:{dates},datasets:[{{label:'Opened',data:{opened},borderColor:'#ffe082',fill:false}},{{label:'Clicked',data:{clicked},borderColor:'#82b1ff',fill:false}},{{label:'Submitted',data:{submitted},borderColor:'#81c784',fill:false}},{{label:'Reported',data:{reported},borderColor:'#ff8a65',fill:false}}]}},options:{{responsive:true,plugins:{{legend:{{display:true}}}}}}}});</script>")
        # Add pagination for summary table
        html.append("<script>function paginateTable(tableId, pageSize){var table=document.getElementById(tableId);var rows=table.getElementsByTagName('tr');var page=1;function showPage(p){for(var i=1;i<rows.length;i++)rows[i].style.display=(i>=(p-1)*pageSize+1&&i<p*pageSize+1)?'':'none';}function createPagination(){var n=Math.ceil((rows.length-1)/pageSize);var pag=document.createElement('div');pag.className='pagination';for(var i=1;i<=n;i++){var btn=document.createElement('button');btn.innerText=i;btn.onclick=(function(i){return function(){page=i;showPage(page);}})(i);pag.appendChild(btn);}table.parentNode.insertBefore(pag,table.nextSibling);}showPage(page);createPagination();}window.onload=function(){paginateTable('summary-table-main',10);};</script>")
        # Add progress bars and badges to recipient details
        html.append("<style>.progress-bar{height:12px;background:#eee;border-radius:6px;overflow:hidden;margin:0.5em 0;}.progress{height:100%;transition:width 0.5s;}.progress-opened{background:#ffe082;}.progress-clicked{background:#82b1ff;}.progress-submitted{background:#81c784;}.progress-reported{background:#ff8a65;}</style>")
        # Add actionable insights section
        html.append("<div class='accordion-section' id='insights'><div class='accordion-header' onclick=\"toggleAccordion('insights')\">Actionable Insights</div><div class='accordion-content'><ul>")
        if pattern_stats['weak']>0:
            html.append(f"<li><b>{pattern_stats['weak']}</b> weak passwords detected. Recommend user training and password policy enforcement.</li>")
        if pattern_stats['reused']>0:
            html.append(f"<li><b>{pattern_stats['reused']}</b> reused passwords detected. Recommend password reuse monitoring.</li>")
        if total_clicked>0 and total_reported/total_clicked<0.5:
            html.append("<li>Reporting rate is low compared to click rate. Recommend awareness campaigns.</li>")
        if sum(opened)>0 and sum(clicked)/sum(opened)<0.2:
            html.append("<li>Click rate is low compared to open rate. Email content may not be enticing enough or users are cautious.</li>")
        if sum(submitted)>0 and sum(submitted)/sum(clicked)<0.2:
            html.append("<li>Submission rate is low compared to click rate. Landing page may not be convincing or users are cautious.</li>")
        if len(device_counter)>1:
            html.append("<li>Multiple device types detected. Ensure security awareness covers all device types.</li>")
        if len(geo_counter)>1:
            html.append("<li>Multiple locations detected. Consider geo-targeted training or investigation.</li>")
        html.append("</ul></div></div>")
        html.append("</body></html>")
        with open(self.output_html_report, "w", encoding="utf-8") as f:
            f.write("\n".join(html))
        print(f"[Dedsec Technologies] Done! Check '{self.output_html_report}' for your HTML report.")

    def config_section_map(self, config_parser, section):
        """This function helps by reading accepting a config file section, from gophish.config,
        and returning a dictionary object that can be referenced for configuration settings.
        """
        section_dict = {}
        options = config_parser.options(section)
        for option in options:
            try:
                section_dict[option] = config_parser.get(section, option)
                if section_dict[option] == -1:
                    print(f"[-] Skipping: {option}")
            except:
                print(f"[!] There was an error with: {option}")
                section_dict[option] = None
        return section_dict

    # Add color mapping utility
    def get_action_color(self, opened, clicked, submitted, reported):
        """Return a color code (Excel hex or Word RGB tuple) based on action criticality and overlays for multiple actions."""
        # Priority: submit > click > open > report > none
        # Maroon for submit, blue for click, yellow for open, green for report, grey for none
        # Overlay: if multiple actions, blend or use a pattern (for now, use a striped/combined color or a special color)
        actions = []
        if submitted:
            actions.append('submitted')
        if clicked:
            actions.append('clicked')
        if opened:
            actions.append('opened')
        if reported:
            actions.append('reported')
        if len(actions) == 0:
            return {'excel': '#B0B0B0', 'word': (0xB0, 0xB0, 0xB0), 'desc': 'none'}  # grey
        if len(actions) == 1:
            if actions[0] == 'submitted':
                return {'excel': '#800000', 'word': (0x80, 0x00, 0x00), 'desc': 'submit'}  # maroon
            if actions[0] == 'clicked':
                return {'excel': '#0070C0', 'word': (0x00, 0x70, 0xC0), 'desc': 'click'}  # blue
            if actions[0] == 'opened':
                return {'excel': '#FFD700', 'word': (0xFF, 0xD7, 0x00), 'desc': 'open'}  # yellow
            if actions[0] == 'reported':
                return {'excel': '#00B050', 'word': (0x00, 0xB0, 0x50), 'desc': 'report'}  # green
        # Multiple actions: use a special overlay color (purple for submit+click, orange for open+click, etc.)
        if 'submitted' in actions:
            if 'clicked' in actions:
                return {'excel': '#800080', 'word': (0x80, 0x00, 0x80), 'desc': 'submit+click'}  # purple
            return {'excel': '#800000', 'word': (0x80, 0x00, 0x00), 'desc': 'submit'}
        if 'clicked' in actions and 'opened' in actions:
            return {'excel': '#FF8C00', 'word': (0xFF, 0x8C, 0x00), 'desc': 'open+click'}  # orange
        if 'clicked' in actions:
            return {'excel': '#0070C0', 'word': (0x00, 0x70, 0xC0), 'desc': 'click'}
        if 'opened' in actions:
            return {'excel': '#FFD700', 'word': (0xFF, 0xD7, 0x00), 'desc': 'open'}
        if 'reported' in actions:
            return {'excel': '#00B050', 'word': (0x00, 0xB0, 0x50), 'desc': 'report'}
        return {'excel': '#B0B0B0', 'word': (0xB0, 0xB0, 0xB0), 'desc': 'none'}
